from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

class pyword:
    def __init__(self):
        pass

    def create_doc(self):
        doc = Document()
        doc.add_heading("This is a heading!")   #文档标题
        doc.add_heading("This is a subtitle!", 1)
        doc.add_paragraph("There are two ways, one is through cmd,"
                          "the other is pycharm")
        doc.add_heading("First: Open pycharm", 2)

        #段落中增加文字、并设置字体字号
        run = doc.add_paragraph("Caution: Font-size-20: ").add_run("Caution: Font-size-20")
        run.font.size = Pt(20)

        #设置英文字体
        run = doc.add_paragraph("Here to set English font: ").add_run("hello world")
        run.font.name = "Times new Roman"

        #设置中文字体
        run = doc.add_paragraph("Here to set Chinese font: ").add_run("你好，中国")
        run.font.name = "宋体"

        #设置斜体
        run = doc.add_paragraph("Here to set Italic: ").add_run("Set to italic")
        run.italic = True

        #设置粗体
        run = doc.add_paragraph("Here to set Bold: ").add_run("Set to Bold")
        run.bold = True

        #设置下划线
        run = doc.add_paragraph("Here to set Underline: ").add_run("Set to Underline")
        run.underline = True

        #设置字体颜色
        run = doc.add_paragraph("Here to set Color: ").add_run("Set color")
        run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)

        #增加引用
        run = doc.add_paragraph("Here to set a quote: ", style="Intense Quote")

        #增加无序、有序列表
        doc.add_paragraph("Apple", style="List Bullet")
        doc.add_paragraph("Banan", style="List Bullet")
        doc.add_paragraph("Orange", style="List Bullet")
        doc.add_paragraph("Pear", style="List Bullet")
        doc.add_paragraph("Strawberry", style="List Bullet")
        doc.add_paragraph("Tomato", style="List Bullet")

        doc.add_paragraph().add_run("2021 practice plan!").bold = True
        doc.add_paragraph("Yoga on monday", style="List Number")
        doc.add_paragraph("Running 5 miles on tuesday", style="List Number")
        doc.add_paragraph("Swimming 30 miniutes on wednesday", style="List Number")
        doc.add_paragraph("Climbing 30 miniutes on thursday", style="List Number")
        doc.add_paragraph("Reading 30 miniutes on friday", style="List Number")
        doc.add_paragraph("Table tennis 30 miniutes on saturday", style="List Number")
        doc.add_paragraph("Football 30 miniutes on sunday", style="List Number")

        #图片
        #doc.add_picture("a.png")
        #doc.add_picture("a.png", width=Inches(5), height=Inches(10))

        #表格，包含表头和内容
        table = doc.add_table(rows=3, cols=3)
        for i in range(3):
            for j in range(3):
                hdr_cells = table.rows[i].cells
                hdr_cells[j].text = str(i + j)

        #保存文件
        doc.save("python_word.docx")


        pass

if __name__ == '__main__':
    solution = pyword()
    solution.create_doc()